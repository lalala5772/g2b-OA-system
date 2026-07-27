"""Fills a real-world .docx form (참가신청서 등) with company info, without
requiring the document to contain `{{key}}` placeholders ahead of time.

Approach: dump the document's paragraphs and table cells (with their
indices and exact current text), ask Claude to find the spots that need
company info and what to put there, then apply each instruction only if
the target's current text still matches what Claude was shown — this
guards against index drift or a hallucinated location silently corrupting
an unrelated part of the document. Anything Claude isn't confident about
is simply omitted from its response, which means it's left untouched.
"""

import io

import docx

from app.services import llm_client

MAX_LOCATIONS = 40


def autofill(docx_bytes: bytes, company_text: str) -> tuple[bytes, dict[str, str]]:
    document = docx.Document(io.BytesIO(docx_bytes))
    paragraphs, tables = _dump_structure(document)
    instructions = _plan(paragraphs, tables, company_text)

    filled: dict[str, str] = {}
    for instr in instructions:
        if not isinstance(instr, dict):
            continue
        if instr.get("type") == "paragraph":
            _apply_paragraph(document, instr, filled)
        elif instr.get("type") == "cell":
            _apply_cell(document, instr, filled)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue(), filled


def _dump_structure(document) -> tuple[list[dict], list[dict]]:
    paragraphs = [{"index": i, "text": p.text} for i, p in enumerate(document.paragraphs) if p.text.strip()]
    tables = []
    for t_idx, table in enumerate(document.tables):
        rows = []
        for r_idx, row in enumerate(table.rows):
            cells = [{"col": c_idx, "text": cell.text} for c_idx, cell in enumerate(row.cells)]
            rows.append({"row": r_idx, "cells": cells})
        tables.append({"table": t_idx, "rows": rows})
    return paragraphs, tables


def _plan(paragraphs: list[dict], tables: list[dict], company_text: str) -> list[dict]:
    if not company_text.strip():
        return []

    result = llm_client.ask_json(
        system_prompt=(
            "너는 서류 양식(참가신청서 등)에 회사 정보를 채워주는 어시스턴트다. "
            "주어진 문서 구조(문단과 표)에서 회사 정보(대표자명, 사업자등록번호, 소재지, 상호, "
            "설립일, 자본금, 전화번호, 이메일 등)가 들어가야 할 빈칸을 찾아라. "
            "값을 회사 정보에서 확실히 찾을 수 없으면 그 위치는 결과에서 아예 빼라 — "
            "추측해서 채우지 말고 비워둔다. 원래 라벨/문구는 절대 지우거나 바꾸지 말고, "
            "빈칸에 값만 추가하거나 빈 셀만 채워라. "
            "반드시 JSON 배열만 응답해라. 각 항목은 다음 형태 중 하나:\n"
            '{"type":"paragraph","index":문단번호,"label":"이 항목이 무엇인지",'
            '"value":"채운 값","before":"원래 문단 전체 텍스트","after":"값이 채워진 문단 전체 텍스트"}\n'
            '{"type":"cell","table":표번호,"row":행번호,"col":열번호,"label":"이 항목이 무엇인지",'
            '"value":"채운 값","before":"원래 셀 텍스트","after":"채워진 셀 텍스트"}'
        ),
        user_prompt=(
            f"회사 정보:\n{company_text}\n\n"
            f"문서 문단 목록: {paragraphs}\n\n"
            f"문서 표 목록: {tables}"
        ),
        max_tokens=3000,
    )
    if not isinstance(result, list):
        return []
    return result[:MAX_LOCATIONS]


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def _apply_paragraph(document, instr: dict, filled: dict[str, str]) -> None:
    try:
        index = int(instr["index"])
        before, after = str(instr["before"]), str(instr["after"])
        label, value = str(instr.get("label", "")), str(instr.get("value", ""))
    except (KeyError, TypeError, ValueError):
        return
    if index < 0 or index >= len(document.paragraphs) or before == after:
        return

    paragraph = document.paragraphs[index]
    if paragraph.text != before:
        return

    _replace_paragraph_text(paragraph, after)
    filled[label or before.strip()] = value or after.strip()


def _apply_cell(document, instr: dict, filled: dict[str, str]) -> None:
    try:
        t_idx, r_idx, c_idx = int(instr["table"]), int(instr["row"]), int(instr["col"])
        before, after = str(instr["before"]), str(instr["after"])
        label, value = str(instr.get("label", "")), str(instr.get("value", ""))
    except (KeyError, TypeError, ValueError):
        return
    if before == after or t_idx < 0 or t_idx >= len(document.tables):
        return

    table = document.tables[t_idx]
    if r_idx < 0 or r_idx >= len(table.rows):
        return
    row = table.rows[r_idx]
    if c_idx < 0 or c_idx >= len(row.cells):
        return

    cell = row.cells[c_idx]
    if cell.text != before or not cell.paragraphs:
        return

    _replace_paragraph_text(cell.paragraphs[0], after)
    for extra in cell.paragraphs[1:]:
        _replace_paragraph_text(extra, "")
    filled[label or f"표{t_idx + 1} {r_idx + 1}행 {c_idx + 1}열"] = value or after.strip()
