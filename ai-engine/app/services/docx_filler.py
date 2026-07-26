"""Fills `{{key}}` placeholders in a .docx template.

Not a full templating engine: a paragraph containing a placeholder has its
runs collapsed into one after substitution, so per-run formatting *inside*
that paragraph isn't preserved (paragraph-level style is). Good enough for
"perfect formatting isn't the goal, correct text is" — the same tradeoff
docs/DESIGN.md makes for .hwp parsing.
"""

import io
import re

import docx

PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")


def fill(template_bytes: bytes, field_values: dict[str, str]) -> bytes:
    document = docx.Document(io.BytesIO(template_bytes))

    for paragraph in document.paragraphs:
        _fill_paragraph(paragraph, field_values)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _fill_paragraph(paragraph, field_values)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _fill_paragraph(paragraph, field_values: dict[str, str]) -> None:
    text = paragraph.text
    if "{{" not in text:
        return

    def _replace(match: re.Match) -> str:
        key = match.group(1).strip()
        return field_values.get(key, match.group(0))

    new_text = PLACEHOLDER_PATTERN.sub(_replace, text)
    if new_text == text:
        return

    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)
